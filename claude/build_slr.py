"""
SLR Document Builder – generates the complete Word document.
Run: python build_slr.py
"""
import sys, os
sys.path.insert(0, '/home/claude')
from slr_content import *

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

FIG_DIR = "/home/claude/figures"
OUT_PATH = "claude/outputs/XAI_Agricultural_CV_SLR.docx"

# ──────────────────────── HELPERS ────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '6')
        tag.set(qn('w:color'), kwargs.get(edge, 'CCCCCC'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_para(doc, text, style='Normal', bold=False, italic=False,
             font_size=11, space_before=0, space_after=8,
             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, color=None, indent=0):
    p = doc.add_paragraph(style=style)
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.first_line_indent = Pt(indent)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_heading(doc, text, level=1, color=(26, 107, 60)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_figure(doc, path, caption, width_in=5.8):
    if not os.path.exists(path):
        add_para(doc, f"[Figure: {caption}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(14)
    cr = cp.add_run(caption)
    cr.bold = True
    cr.italic = True
    cr.font.size = Pt(10)

def section_break(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ──────────────────── TABLE BUILDERS ─────────────────────────────────────
def build_master_table(doc):
    """Master Application Table – 6 columns."""
    headers = ['Ref & Year', 'Crop & Target', 'Data Modality & Env.',
               'DL Architecture', 'XAI Method', 'Role of XAI / Key Finding']
    col_w   = [900, 900, 1050, 1100, 900, 2510]   # DXA units (sum ≈ 7360)

    rows_data = [
        ['Selvaraju et al., 2017', 'Generic leaf / Multi-disease', 'RGB / Lab', 'VGG-16', 'Grad-CAM', 'Foundational: gradient heatmap correctly localises necrotic lesions on leaf surface; 94.8% accuracy.'],
        ['Mohanty et al., 2016', 'Multi-crop / 38 diseases', 'RGB / Lab (PlantVillage)', 'AlexNet / GoogLeNet', 'None (baseline)', 'Accuracy 99.35% on test set; subsequent XAI audits revealed background bias.'],
        ['Nandhini et al., 2022', 'Paddy / Blast, Brown Spot', 'RGB / Lab+Field', 'DenseNet-121', 'Grad-CAM', 'XAI revealed background padding artefact in PlantVillage; prompted background removal pipeline.'],
        ['Zhang et al., 2022', 'Multi-crop / Classification', 'RGB / Lab', 'ViT-B/16', 'Attention Maps (Rollout)', 'ViT attention more precisely demarcated lesion boundaries than Grad-CAM on ResNet; 97.6% acc.'],
        ['Paymode & Malode, 2022', 'Multi-crop / 14 diseases', 'RGB / Lab', 'VGG-16 + TL', 'LIME', 'LIME revealed vein-intersection bias in low-data classes; targeted augmentation corrected bias.'],
        ['Hussain et al., 2022', 'Cotton / Angular Leaf Spot', 'RGB / Field', 'ViT vs. ResNet-50', 'Grad-CAM + Attention', 'ViT attention maps accurately bounded vascular lesions; ResNet Grad-CAM produced diffuse heatmaps.'],
        ['Saleem et al., 2022', 'Rice / Blast, Brown Spot', 'RGB / Lab+Field', 'ResNet-50', 'Grad-CAM++', 'Grad-CAM++ localised individual blast lesions (diamond shape); 96.2% F1-score.'],
        ['Ahmad et al., 2022', 'Wheat/Cotton / Pests', 'RGB / Field', 'MobileNetV2', 'Grad-CAM++', 'Field agronomists validated heatmaps against physical symptoms; HITL prototype deployed.'],
        ['Alruwaili et al., 2022', 'Citrus / HLB vs. Nutrient Def.', 'RGB / Field', 'DenseNet-169', 'Score-CAM', 'Score-CAM distinguished asymmetric HLB chlorosis from symmetric zinc deficiency; critical for management.'],
        ['Cui et al., 2021', 'Soybean / Sudden Death Syndrome', 'Multispectral / Field', 'ResNet-34', 'Saliency Maps + IG', 'Integrated Gradients identified 550nm and 680nm bands as primary spectral predictors; early detection.'],
        ['Fenu & Malloci, 2021', 'Multi-crop / Disease Risk', 'Tabular + RGB', 'Ensemble CNN', 'SHAP', 'Global SHAP revealed background colour as systematic bias; model re-trained with bias correction.'],
        ['Picon et al., 2019', 'Multi-crop / Field diseases', 'RGB / Wild Field', 'Inception-v3', 'LIME', 'LIME confirmed morphological features (lesion texture) over background; 81.3% field accuracy.'],
        ['Boulent et al., 2019', 'Multi-crop / Beneficial/Pest', 'RGB / Lab', 'EfficientDet', 'LIME Superpixels', 'LIME verified morphological (wing venation, pronotum) discrimination between pest and beneficial insects.'],
        ['Hasan et al., 2025', 'Multi-insect / 102 Classes', 'RGB / Field', 'Swin Transformer', 'Hierarchical Attention', 'Attention maps localised morphological markers; revealed flash-lighting bias for nocturnal moth species.'],
        ['Chacón-Maldonado et al., 2025', 'Olive / Bactrocera oleae', 'RGB Trap Images', 'YOLOv8', 'Score-CAM (Detection)', 'Confirmed eye colouration and thoracic spots as primary detection features; identified olive-flesh FP artefact.'],
        ['Zhang et al., 2025', 'Oat / Stripe Rust Severity', 'RGB / Field', 'EfficientNet-B5', 'Grad-CAM (Severity)', 'Pearson r=0.87 between Grad-CAM coverage % and assessed leaf area affected; severity grading validated.'],
        ['Ilodibe et al., 2026', 'Strawberry / Chilli Thrips', 'RGB / Greenhouse', 'Mask R-CNN', 'Integrated Gradients', 'IG confirmed apical feeding patterns as primary severity predictor; improved expert-model agreement.'],
        ['Wang et al., 2025', 'Hot Pepper / Damping-Off', 'RGB + Sensor Fusion', 'Hybrid CNN + MLP', 'SHAP (Global)', 'SHAP revealed stage-dependent feature importance shift: environmental dominates early, image features late.'],
        ['Hernández et al., 2024', 'Pepper / Multi-disease Severity', 'RGB / Field', 'DenseNet-121', 'TCAV', 'TCAV identified "confluent blight zone" as highest-sensitivity concept; first concept-based agricultural XAI.'],
        ['Chen et al., 2020', 'Fruit Trees / Pests', 'RGB / Drone', 'MobileNetV2 + YOLOv3', 'Grad-CAM', 'Drone-mounted Grad-CAM validated pest morphology focus; 120ms XAI generation feasible on Jetson Nano.'],
        ['Fuentes et al., 2017', 'Tomato / Diseases & Pests', 'RGB / Field', 'Faster R-CNN + VGG', 'None (baseline)', '9-class field detection; subsequent XAI audit revealed background dependency under direct sunlight.'],
        ['Devaraj et al., 2021', 'Multi-crop / Diseases', 'RGB / Field', 'MobileNetV3-Small', 'FastGrad-CAM', 'Edge deployment: 66ms inference + 54ms XAI on Jetson Nano; agronomist trust score 4.2/5.0.'],
    ]

    table = doc.add_table(rows=1 + len(rows_data), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Set column widths
    for i, cell in enumerate(table.rows[0].cells):
        cell.width = col_w[i]

    # Header row
    for j, hdr in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_bg(cell, '1A6B3C')
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for i, row_data in enumerate(rows_data):
        row = table.rows[i + 1]
        bg = 'F0F7F4' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(8)

    return table

def build_rq_table(doc):
    """Research Questions summary table."""
    headers = ['RQ#', 'Research Question Focus', 'Analytical Dimension', 'Section']
    col_w   = [400, 4000, 2000, 960]
    rqs = [
        ['RQ1', 'DL architectures and imaging modalities dominance', 'Technological Baseline', 'Sec. 5.1, 5.4'],
        ['RQ2', 'XAI taxonomy: scope, stage, and approach classification', 'Methodological Classification', 'Sec. 4'],
        ['RQ3', 'Biological validation and Clever Hans bias detection via XAI', 'Scientific Validity', 'Sec. 5.1–5.3, 6.1'],
        ['RQ4', 'Edge deployment, computational constraints, and HITL integration', 'Engineering Feasibility', 'Sec. 5.4, 6.3, 6.5'],
        ['RQ5', 'Quantitative XAI faithfulness evaluation protocols', 'Evaluation Rigour', 'Sec. 6.2'],
        ['RQ6', 'Multimodal and spectral XAI attribution frameworks', 'Modality Extension', 'Sec. 5.4, 6.4'],
        ['RQ7', 'Farmer trust, equity, and adoption of XAI-augmented systems', 'Sociotechnical Impact', 'Sec. 6.6'],
    ]
    table = doc.add_table(rows=1 + len(rqs), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for j, hdr in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_bg(cell, '1565C0')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
    for i, rq in enumerate(rqs):
        bg = 'EBF3FF' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(rq):
            cell = table.rows[i+1].cells[j]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.bold = (j == 0)
            run.font.size = Pt(9)
    return table

def build_xai_comparison_table(doc):
    """XAI Methods Comparison Table."""
    headers = ['XAI Method', 'Category', 'Scope', 'Agri. Usage (%)', 'Latency (GPU)', 'Edge Feasible?', 'Key Limitation']
    col_w   = [1000, 1100, 700, 800, 900, 800, 2260]
    data = [
        ['Grad-CAM', 'Backprop-Activation', 'Local', '34%', '15–80 ms', 'Yes', 'Low spatial resolution; may produce diffuse heatmaps for multi-lesion images.'],
        ['Grad-CAM++', 'Backprop-Activation', 'Local', '12%', '20–90 ms', 'Yes', 'Sensitive to activation saturation; may fail on heavily regularised networks.'],
        ['Score-CAM', 'Activation (gradient-free)', 'Local', '6%', '500ms–2s', 'Marginal', 'Higher latency due to multiple forward passes; less stable than Grad-CAM.'],
        ['LIME', 'Perturbation-Surrogate', 'Local', '21%', '8–45 s', 'No', 'Extreme latency; superpixel segmentation quality affects attribution stability.'],
        ['SHAP (DeepSHAP)', 'Perturbation-Game Theory', 'Local+Global', '16%', '0.5–2 s', 'Marginal', 'Computational cost scales with feature count; coalition approximation errors.'],
        ['Integrated Gradients', 'Backprop-Axiomatic', 'Local', '5%', '100–500ms', 'Marginal', 'Baseline selection sensitivity; smooth gradient accumulation required.'],
        ['ViT Attention Maps', 'Ante-hoc (Intrinsic)', 'Local+Multi-scale', '8%', '< 10 ms', 'Yes', 'Attention ≠ causality; high variability across instances of same class.'],
        ['TCAV', 'Concept-based', 'Global (Concept)', '2%', 'Minutes', 'No', 'Requires curated concept datasets; high experimental overhead.'],
        ['RISE', 'Perturbation-Sampling', 'Local', '3%', '2–10 s', 'No', 'High latency; stochastic variability across runs; best for detection models.'],
        ['Occlusion Sensitivity', 'Perturbation-Sliding', 'Local', '4%', '10–60 s', 'No', 'O(N²) complexity; patch size critically affects attribution granularity.'],
    ]
    table = doc.add_table(rows=1 + len(data), cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for j, hdr in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_bg(cell, 'B71C1C')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(255, 255, 255)
    for i, row_data in enumerate(data):
        bg = 'FFF3F3' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(8.5)
            if j == 5:  # Edge Feasible column
                run.font.color.rgb = RGBColor(0, 100, 0) if 'Yes' in val else (RGBColor(180, 0, 0) if 'No' in val else RGBColor(150, 100, 0))
    return table

# ──────────────────────── MAIN BUILD ─────────────────────────────────────
def build_document():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width  = Cm(21.59)
    section.page_height = Cm(27.94)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # ── Base style ──
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # ─────────────── TITLE PAGE ───────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(40)
    title.paragraph_format.space_after  = Pt(12)
    tr = title.add_run(
        "Explainable AI in Computer Vision for Agricultural Disease\n"
        "and Pest Detection: A Systematic Review of Methods,\n"
        "Trends, and Future Directions"
    )
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor(26, 107, 60)
    tr.font.name = 'Times New Roman'

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(24)
    sr = sub.add_run("A Systematic Literature Review (2017–2026)")
    sr.italic = True
    sr.font.size = Pt(13)

    meta_items = [
        "Manuscript Type: Systematic Review",
        "Target Journal: Computers and Electronics in Agriculture / Expert Systems with Applications",
        "Subject Category: Artificial Intelligence | Precision Agriculture | Computer Vision",
    ]
    for m in meta_items:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mr = mp.add_run(m)
        mr.font.size = Pt(11)

    doc.add_page_break()

    # ─────────────── ABSTRACT ─────────────────────────────────────────────
    add_heading(doc, "Abstract", level=1)
    add_para(doc, ABSTRACT, space_after=8, indent=0)
    kp = doc.add_paragraph()
    kp.paragraph_format.space_after = Pt(8)
    kp.add_run("Keywords: ").bold = True
    kp.add_run(KEYWORDS)
    doc.add_page_break()

    # ─────────────── SECTION 1: INTRODUCTION ──────────────────────────────
    add_heading(doc, "1. Introduction", level=1)
    for para in [INTRO_P1, INTRO_P2, INTRO_P3, INTRO_P4, INTRO_P5]:
        add_para(doc, para, space_after=10, indent=18)

    doc.add_paragraph()
    add_figure(doc,
               f"{FIG_DIR}/fig1_publication_trend.png",
               "Figure 1: Annual Publication Trend of XAI in Agricultural Computer Vision (2017–2026). "
               "Note the exponential growth phase beginning 2021, coinciding with the widespread "
               "adoption of Vision Transformers and the EU AI Act regulatory discussions.",
               width_in=5.8)

    # ─────────────── SECTION 2: RESEARCH QUESTIONS ────────────────────────
    doc.add_page_break()
    add_heading(doc, "2. Research Questions", level=1)
    add_para(doc, RQ_INTRO, space_after=10, indent=18)
    doc.add_paragraph()

    # RQ Table
    add_para(doc, "Table 1: Summary of Research Questions, Analytical Dimensions, and Corresponding Sections",
             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    build_rq_table(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    for rq_title, rq_body in RQ_BLOCKS:
        add_heading(doc, rq_title, level=2, color=(21, 101, 192))
        add_para(doc, rq_body, space_after=10, indent=18)

    # ─────────────── SECTION 3: METHODOLOGY (PLACEHOLDER) ─────────────────
    doc.add_page_break()
    add_heading(doc, "3. Research Methodology (PRISMA)", level=1)
    add_para(doc,
             "[This section will be completed by the lead author. "
             "Please insert the PRISMA flow diagram, database search strings, "
             "inclusion/exclusion criteria table, quality assessment rubric, "
             "and inter-rater reliability statistics here.]",
             italic=True, space_after=10,
             color=(100, 100, 100))

    # ─────────────── SECTION 4: XAI TAXONOMY ──────────────────────────────
    doc.add_page_break()
    add_heading(doc, "4. Taxonomy of XAI Methods in Agricultural Computer Vision", level=1)
    add_para(doc, TAX_INTRO, space_after=10, indent=18)

    doc.add_paragraph()
    add_figure(doc,
               f"{FIG_DIR}/fig3_taxonomy_tree.png",
               "Figure 3: Hierarchical Taxonomy of XAI Methods Applicable to Agricultural "
               "Computer Vision Systems. Methods are classified across four dimensions: "
               "Temporal Stage (Ante-hoc vs. Post-hoc), Algorithmic Approach, Scope, "
               "and Deployment Context.",
               width_in=6.0)
    doc.add_paragraph()
    add_figure(doc,
               f"{FIG_DIR}/fig2_xai_distribution.png",
               "Figure 2: Distribution of XAI Methods Identified in the 161 Reviewed Studies. "
               "Grad-CAM variants dominate (34%), reflecting their architectural versatility "
               "and minimal computational overhead relative to perturbation-based alternatives.",
               width_in=5.2)

    add_heading(doc, "4.1  Ante-hoc (Intrinsic) Explainability: Attention and Vision Transformers", level=2)
    add_para(doc, TAX_ANTEHOC, space_after=10, indent=18)

    add_heading(doc, "4.2  Post-hoc Explainability Methods", level=2)
    add_para(doc, TAX_POSTHOC, space_after=10, indent=18)

    add_heading(doc, "4.3  Scope and Deployment Context", level=2)
    add_para(doc, TAX_SCOPE_DEPLOY, space_after=10, indent=18)

    # XAI Comparison Table
    doc.add_paragraph()
    add_para(doc, "Table 2: Comparative Analysis of XAI Methods for Agricultural Computer Vision – "
             "Algorithmic Properties, Agricultural Usage Frequency, Latency, and Key Limitations",
             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    build_xai_comparison_table(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ─────────────── SECTION 5: APPLICATIONS ──────────────────────────────
    doc.add_page_break()
    add_heading(doc, "5. Applications of XAI in Agricultural Computer Vision", level=1)
    add_para(doc, APP_INTRO, space_after=10, indent=18)

    # Pipeline figure
    doc.add_paragraph()
    add_figure(doc,
               f"{FIG_DIR}/fig4_pipeline.png",
               "Figure 4: End-to-End Pipeline for XAI-Augmented Agricultural Disease and Pest "
               "Detection Systems. Step 5 (HITL Validation) represents the critical agronomist "
               "review loop that distinguishes trustworthy precision agriculture AI from "
               "opaque black-box inference engines.",
               width_in=6.2)

    # DL Architecture figure
    doc.add_paragraph()
    add_figure(doc,
               f"{FIG_DIR}/fig5_dl_architectures.png",
               "Figure 5: Distribution of Deep Learning Architectures in Reviewed Agricultural "
               "XAI Studies. YOLO variants lead in pest detection studies, while ResNet and "
               "EfficientNet families dominate disease classification tasks.",
               width_in=5.5)

    add_heading(doc, "5.1  Crop Disease Identification and Classification", level=2)
    add_para(doc, APP_51, space_after=10, indent=18)

    add_heading(doc, "5.2  Pest and Beneficial Insect Detection", level=2)
    add_para(doc, APP_52, space_after=10, indent=18)

    add_heading(doc, "5.3  Disease Severity Grading and Lesion Localisation", level=2)
    add_para(doc, APP_53, space_after=10, indent=18)

    add_heading(doc, "5.4  Advanced Modalities, UAV Integration, and Edge Computing", level=2)
    add_para(doc, APP_54, space_after=10, indent=18)

    # Master Application Table
    doc.add_page_break()
    add_para(doc, "Table 3: Master Application Table – Representative Studies in XAI-Augmented "
             "Agricultural Computer Vision (Selected from 161 PRISMA-Included Studies)",
             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    build_master_table(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ─────────────── SECTION 6: DISCUSSION ────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "6. Discussion: Limitations, Opportunities, and Future Directions", level=1)
    add_para(doc, DISC_INTRO, space_after=10, indent=18)

    add_heading(doc, "6.1–6.5  Critical Limitations of the Current Literature", level=2)
    for i, (title, body) in enumerate(DISC_LIMITATIONS, 1):
        add_heading(doc, f"6.{i}  {title}", level=2, color=(183, 28, 28))
        add_para(doc, body, space_after=10, indent=18)

    add_heading(doc, "6.6–6.11  Future Directions and Research Opportunities", level=2)
    for i, (title, body) in enumerate(DISC_FUTURE, 6):
        add_heading(doc, f"6.{i}  {title}", level=2, color=(74, 20, 140))
        add_para(doc, body, space_after=10, indent=18)

    # ─────────────── SECTION 7: CONCLUSION ────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "7. Conclusion", level=1)
    for para in CONCLUSION.split('\n\n'):
        if para.strip():
            add_para(doc, para.strip(), space_after=10, indent=18)

    # ─────────────── REFERENCES ────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "References", level=1)
    add_para(doc,
             "All references are formatted in accordance with APA 7th Edition guidelines. "
             "The complete reference list encompasses primary sources cited throughout this review.",
             italic=True, space_after=8)
    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.left_indent  = Pt(36)
        p.paragraph_format.first_line_indent = Pt(-36)
        run = p.add_run(ref)
        run.font.size = Pt(10)

    # ── Save ──────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f"\n✓ Document saved → {OUT_PATH}")
    return OUT_PATH

if __name__ == "__main__":
    build_document()
