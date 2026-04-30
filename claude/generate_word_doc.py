"""
Generate Word Document from SLR Content
"""
import sys
from pathlib import Path

# Add the directory to path to import slr_content
sys.path.insert(0, str(Path(__file__).parent))

from slr_content import (
    ABSTRACT, KEYWORDS, INTRO_P1, INTRO_P2, INTRO_P3, INTRO_P4, INTRO_P5,
    RQ_INTRO, RQ_BLOCKS, TAX_INTRO, TAX_ANTEHOC, TAX_POSTHOC, TAX_SCOPE_DEPLOY,
    APP_INTRO, APP_51, APP_52, APP_53, APP_54,
    DISC_INTRO, DISC_LIMITATIONS, DISC_FUTURE, CONCLUSION, REFERENCES
)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a Document object
doc = Document()

# Set up document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Helper functions
def add_title(text):
    """Add title to document"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.font.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return p

def add_heading(text, level=1):
    """Add heading"""
    p = doc.add_heading(text, level=level)
    return p

def add_text(text):
    """Add paragraph text"""
    if text:
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing = 1.15
        return p

def add_numbered_items(items, start_num=1):
    """Add numbered list items"""
    for i, item in enumerate(items, start=start_num):
        p = doc.add_paragraph(style='List Number')
        p.text = item
        p.paragraph_format.line_spacing = 1.15

# Title
add_title("Systematic Literature Review")
add_title("Explainable AI in Agricultural Computer Vision")
add_text("Pest and Disease Detection (2017-2026)")
doc.add_paragraph()

# Abstract
add_heading("Abstract", level=1)
add_text(ABSTRACT)
doc.add_paragraph()

# Keywords
add_heading("Keywords", level=1)
add_text(KEYWORDS)
doc.add_paragraph()

# Section 1: Introduction
add_heading("1. Introduction", level=1)
add_text(INTRO_P1)
add_text(INTRO_P2)
add_text(INTRO_P3)
add_text(INTRO_P4)
add_text(INTRO_P5)
doc.add_paragraph()

# Section 2: Research Questions
add_heading("2. Research Questions", level=1)
add_text(RQ_INTRO)
doc.add_paragraph()
for i, (title, content) in enumerate(RQ_BLOCKS, start=1):
    add_heading(title, level=2)
    add_text(content)
    if i < len(RQ_BLOCKS):
        doc.add_paragraph()

# Section 4: XAI Taxonomy
add_heading("3. XAI Taxonomy", level=1)
add_text(TAX_INTRO)
doc.add_paragraph()

add_heading("3.1 Ante-hoc Interpretability", level=2)
add_text(TAX_ANTEHOC)
doc.add_paragraph()

add_heading("3.2 Post-hoc Explanations", level=2)
add_text(TAX_POSTHOC)
doc.add_paragraph()

add_text(TAX_SCOPE_DEPLOY)
doc.add_paragraph()

# Section 5: Applications
add_heading("4. Application Domains", level=1)
add_text(APP_INTRO)
doc.add_paragraph()

add_heading("4.1 Crop Disease Classification", level=2)
add_text(APP_51)
doc.add_paragraph()

add_heading("4.2 Pest and Arthropod Detection", level=2)
add_text(APP_52)
doc.add_paragraph()

add_heading("4.3 Disease Severity and Lesion Localisation", level=2)
add_text(APP_53)
doc.add_paragraph()

add_heading("4.4 Advanced Modalities, UAV Integration, and Edge Computing", level=2)
add_text(APP_54)
doc.add_paragraph()

# Section 6: Discussion
add_heading("5. Discussion", level=1)
add_text(DISC_INTRO)
doc.add_paragraph()

add_heading("5.1 Critical Limitations", level=2)
for title, content in DISC_LIMITATIONS:
    add_heading(title, level=3)
    add_text(content)
    doc.add_paragraph()

add_heading("5.2 Future Research Directions", level=2)
for title, content in DISC_FUTURE:
    add_heading(title, level=3)
    add_text(content)
    doc.add_paragraph()

# Section 7: Conclusion
add_heading("6. Conclusion", level=1)
add_text(CONCLUSION)
doc.add_paragraph()

# References
add_heading("7. References", level=1)
for ref in REFERENCES:
    p = doc.add_paragraph(ref, style='List Bullet')
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)

# Save the document
output_path = Path(__file__).parent / "XAI_Agricultural_SLR.docx"
doc.save(str(output_path))
print(f"✓ Word document created successfully: {output_path}")
print(f"  File size: {output_path.stat().st_size / 1024:.1f} KB")
